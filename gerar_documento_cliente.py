from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colors from the site brand
PRIMARY = RGBColor(0x1a, 0x17, 0x14)   # dark brown/black
SECONDARY = RGBColor(0x78, 0x57, 0x2b) # warm gold
GREY = RGBColor(0x6b, 0x72, 0x80)      # zinc-500
LIGHT = RGBColor(0xf4, 0xf4, 0xf5)     # zinc-100

def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_heading_paragraph(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = SECONDARY if level == 1 else PRIMARY
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = PRIMARY
    return p

def add_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = GREY
    return p

def add_checkbox_item(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = PRIMARY
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'e4e4e7')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


doc = Document()

# Page margins
section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ─── HEADER BAND ────────────────────────────────────────────────
header_para = doc.add_paragraph()
header_para.paragraph_format.space_before = Pt(0)
header_para.paragraph_format.space_after = Pt(0)
# Gold left border via shading workaround: use a table as header bar
header_table = doc.add_table(rows=1, cols=2)
header_table.autofit = False
header_table.columns[0].width = Cm(0.6)
header_table.columns[1].width = Cm(15.4)

# Gold accent strip
set_cell_background(header_table.rows[0].cells[0], '78572b')
header_table.rows[0].cells[0].paragraphs[0].add_run(' ')

# Title cell
title_cell = header_table.rows[0].cells[1]
set_cell_background(title_cell, '1a1714')
title_p = title_cell.paragraphs[0]
title_p.paragraph_format.left_indent = Cm(0.5)
title_p.paragraph_format.space_before = Pt(10)
title_p.paragraph_format.space_after = Pt(4)
r1 = title_p.add_run('Informações Necessárias para o Website')
r1.bold = True
r1.font.size = Pt(16)
r1.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

sub_p = title_cell.add_paragraph()
sub_p.paragraph_format.left_indent = Cm(0.5)
sub_p.paragraph_format.space_before = Pt(0)
sub_p.paragraph_format.space_after = Pt(10)
r2 = sub_p.add_run('Dra. Conceição Lopes  ·  Website  ·  2025')
r2.font.size = Pt(9)
r2.font.color.rgb = RGBColor(0xa1, 0x8a, 0x6a)

# ─── INTRO ──────────────────────────────────────────────────────
doc.add_paragraph()
intro = doc.add_paragraph()
intro.paragraph_format.space_after = Pt(4)
r = intro.add_run('Olá Dra. Conceição,')
r.bold = True
r.font.size = Pt(10.5)
r.font.color.rgb = PRIMARY

body_intro = doc.add_paragraph(
    'Para podermos finalizar o website, precisamos da sua ajuda com algumas decisões pontuais. '
    'Cada ponto tem impacto direto no que os potenciais clientes vêem.'
)
body_intro.paragraph_format.space_after = Pt(4)
for run in body_intro.runs:
    run.font.size = Pt(10.5)
    run.font.color.rgb = GREY

add_divider(doc)

# ─── PERGUNTA 1 ─────────────────────────────────────────────────
add_heading_paragraph(doc, '1.  A Primeira Consulta é Gratuita?')

add_label(doc, 'O que precisamos saber:')
add_body(doc, 'A consulta inicial tem custo, ou a primeira conversa é gratuita?')

add_label(doc, 'Porquê é importante:')
add_body(doc,
    'Se a primeira consulta for genuinamente gratuita, colocamos isso em destaque no botão '
    'de agendamento, o que aumenta significativamente o número de pessoas que contactam.'
)

add_label(doc, 'A sua resposta:')
add_checkbox_item(doc, 'A primeira consulta é gratuita')
add_checkbox_item(doc, 'A primeira consulta tem um custo fixo (qual? ________________)')
add_checkbox_item(doc, 'O custo depende do caso')

add_divider(doc)

# ─── PERGUNTA 2 ─────────────────────────────────────────────────
add_heading_paragraph(doc, '2.  Avaliações e Testemunhos do Google')

add_body(doc,
    'Sabemos que tem avaliações no Google Business. Temos três opções para as apresentar no website:'
)

# Options table
tbl = doc.add_table(rows=4, cols=3)
tbl.style = 'Table Grid'

headers = ['Opção', 'O que inclui', 'Custo extra']
header_row = tbl.rows[0]
for i, h in enumerate(headers):
    cell = header_row.cells[i]
    set_cell_background(cell, '1a1714')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

rows_data = [
    ('A — Estático à sua escolha',
     'A Dra. Conceição escolhe 3 reviews do Google que aparecem fixas no website',
     '€0'),
    ('B — Estático selecionado por nós',
     'Selecionamos os 3 melhores reviews atuais e implementamos',
     '€0'),
    ('C — Dinâmico\n(recomendado para o futuro)',
     'Rating e reviews sempre atualizados automaticamente via Google Business API. '
     'Requer acesso ao Google Business.',
     '+€60'),
]

for i, (opt, desc, cost) in enumerate(rows_data):
    row = tbl.rows[i + 1]
    bg = 'faf9f7' if i % 2 == 0 else 'f4f4f5'
    for cell in row.cells:
        set_cell_background(cell, bg)

    p0 = row.cells[0].paragraphs[0]
    p0.paragraph_format.space_before = Pt(4)
    p0.paragraph_format.space_after = Pt(4)
    r = p0.add_run(opt)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = SECONDARY

    p1 = row.cells[1].paragraphs[0]
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(desc)
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = PRIMARY

    p2 = row.cells[2].paragraphs[0]
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(4)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(cost)
    r2.bold = True
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = SECONDARY if cost != '€0' else PRIMARY

# Column widths
tbl.columns[0].width = Cm(4.2)
tbl.columns[1].width = Cm(9.5)
tbl.columns[2].width = Cm(2.3)

note_p = doc.add_paragraph()
note_p.paragraph_format.space_before = Pt(8)
note_p.paragraph_format.space_after = Pt(4)
rn = note_p.add_run('Porquê a Opção C é recomendada para o futuro:  ')
rn.bold = True
rn.font.size = Pt(9.5)
rn.font.color.rgb = GREY
rn2 = note_p.add_run(
    'Com poucas avaliações, o número muda raramente e a diferença entre dinâmico e estático é mínima. '
    'A partir das 20 a 25 avaliações, novos reviews começam a chegar com regularidade suficiente para '
    'que o conteúdo estático fique desatualizado em poucos meses. A partir desse ponto, a integração dinâmica '
    'evita ter de pedir manualmente para atualizar o website de cada vez que aparece uma nova avaliação, '
    'e garante que os visitantes veem sempre os números reais e atuais. '
    'Inclui configuração de permissões no Google Business, integração técnica, testes e documentação.'
)
rn2.font.size = Pt(9.5)
rn2.font.color.rgb = GREY

add_label(doc, 'A sua resposta:  Opção ____')

add_divider(doc)

# ─── PERGUNTA 3 ─────────────────────────────────────────────────
add_heading_paragraph(doc, '3.  A Sua História: Porque Escolheu o Direito?')

add_label(doc, 'O que precisamos saber:')
add_body(doc, 'O que a levou a ser advogada? Houve algum momento, experiência ou razão pessoal que a motivou?')

add_label(doc, 'Porquê é importante:')
add_body(doc,
    'A página "Sobre" tem um perfil profissional sólido. Uma breve frase pessoal torna-a '
    'mais próxima para clientes que estão a passar por um momento difícil.'
)

add_label(doc, 'A sua resposta (1 a 2 frases):')
p_blank = doc.add_paragraph()
p_blank.paragraph_format.space_after = Pt(30)

add_divider(doc)

# ─── PERGUNTA 4 ─────────────────────────────────────────────────
add_heading_paragraph(doc, '4.  Disponibilidade Real (Agenda)')

add_label(doc, 'O que precisamos saber:')
add_body(doc, 'A sua agenda está habitualmente preenchida com antecedência, ou tem disponibilidade imediata?')

add_label(doc, 'Porquê é importante:')
add_body(doc,
    'Se a agenda for disputada, usamos isso no website ("Agenda limitada — reserve com antecedência"), '
    'o que cria urgência natural e melhora a perceção de procura pelo seu trabalho.'
)

add_label(doc, 'A sua resposta:')
add_checkbox_item(doc, 'Tenho disponibilidade quase imediata (poucos dias)')
add_checkbox_item(doc, 'A agenda está ocupada com 1 a 2 semanas de antecedência')
add_checkbox_item(doc, 'A agenda está muito preenchida (mais de 2 semanas)')

add_divider(doc)

# ─── FOOTER NOTE ────────────────────────────────────────────────
footer_p = doc.add_paragraph()
footer_p.paragraph_format.space_before = Pt(8)
rf = footer_p.add_run(
    'Assim que nos responder podemos publicar a versão final do website. '
    'Não há urgência — responda quando lhe for conveniente. Obrigado.'
)
rf.font.size = Pt(10)
rf.italic = True
rf.font.color.rgb = GREY

doc.save(r'e:\JamesDev\Projetos Clientes\Advogada\Perguntas_Website_DraConceicaoLopes.docx')
print('Documento criado com sucesso.')
